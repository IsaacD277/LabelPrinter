from docx import Document
from docx.shared import Pt
from datetime import datetime
import win32api
import config
import requests
import os

def initialize_api():
    """Initialize API base URL and headers."""
    base_url = "http://na.myconnectwise.net/v4_6_release/apis/3.0"
    client_id = config.client_id
    login_company = config.login_company
    public_key = config.public_key
    private_key = config.private_key
    auth = (f"{login_company}+{public_key}", private_key)
    headers = {
        "clientID": client_id,
        "Authorization": requests.auth._basic_auth_str(auth[0], auth[1])
    }

    return base_url, headers

def fetch_ticket_and_client(base_url, headers, ticket_id):
    if not ticket_id:
        return {}  # Return an empty dictionary if ticket_id is None or invalid

    url = f"{base_url}/service/tickets/{ticket_id}?fields=id,company/name"

    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        details = response.json()
    except requests.RequestException as e:
        try:
            projectUrl = f"{base_url}/project/tickets/{ticket_id}?fields=id,company/name"
            response = requests.get(projectUrl, headers=headers)
            response.raise_for_status()
            details = response.json()
        except requests.RequestException as projecte:
            return {"error": f"Service and Project fetch failed: {str(e)} | {str(projecte)}"}

    # Extract ticket number and client name from the response
    ticket_number = details.get("id")
    client_name = details.get("company").get("name")

    # If client name is longer than 30 characters, remove the end of the string on a word break
    if len(client_name) > 30:
        splitClientName = client_name[:30].split(" ")
        client_name = splitClientName[0]
        for word in splitClientName[1:-1]:
            client_name += " " + word

    return ticket_number, client_name

def update_Word_document(base_dir, ticket_number, client_name, date, laptop = False):
    output_path = os.path.join(base_dir, "ETS EQUIPMENT SHEET_Filled.docx")  # Output file with filled details

    # Print the Word document with details for laptops
    if laptop: # If the ticket number starts with "L" (for laptop boxes)
        template_path = os.path.join(base_dir, "ETS EQUIPMENT SHEET TEMPLATE - Laptop.docx")
        doc = Document(template_path)

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if "Ticket#" in paragraph.text:
                run = paragraph.add_run(f"{ticket_number}")
                run.bold = True
                run.font.size = Pt(11)

            elif "Date:" in paragraph.text:
                run = paragraph.add_run(f"{date}")
                run.font.size = Pt(11)

            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {client_name}")
                run.bold = True
                run.font.size = Pt(11)

        # Save the modified document
        try:
            doc.save(output_path)
        except PermissionError:
            input(f"Please close document and press Enter to try again.")

    # Print the Word document with details for other boxes
    else: # If the ticket number does not start with "L" (for other boxes)
        template_path = os.path.join(base_dir, "ETS EQUIPMENT SHEET TEMPLATE - Normal.docx")
        doc = Document(template_path)

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if "Ticket #:" in paragraph.text:
                run = paragraph.add_run(f" {ticket_number}")
                run.bold = True
                run.font.size = Pt(11)

                run = paragraph.add_run(f" \tDate: {date}         Tech:______")
                run.font.size = Pt(11)

            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {client_name}")
                run.bold = True
                run.font.size = Pt(11)

        # Save the modified document
        try:
            doc.save(output_path)
        except PermissionError:
            input(f"Please close document and press Enter to try again.")

    return output_path

def main():
    base_url, headers = initialize_api()
    base_dir = os.path.dirname(os.path.abspath(__file__))

    while True:
        # Prompt user for Service Ticket # and exit if they type "end"
        userInput = input("Enter Service Ticket # (Start with \'L\' for laptops): ")
        ticketNumberInput = userInput.replace("L", "").strip()

        if ticketNumberInput.lower() == "end":
            break

        # Set variables for Ticket#, Date, and Client Name
        ticket_number, client_name = fetch_ticket_and_client(base_url=base_url, headers=headers, ticket_id=ticketNumberInput)
        date = datetime.now().strftime("%m/%d/%y")

        # Update Word document with ticket details   
        output_path = update_Word_document(base_dir, ticket_number, client_name, date, laptop=userInput[0] == "L")
        
        # Print to label printer
        win32api.ShellExecute(0, "printto", output_path, f'"{config.printer_name}"', ".", 0)

if __name__ == "__main__":
    main()