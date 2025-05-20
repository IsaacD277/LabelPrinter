from docx import Document
from docx.shared import Pt
from datetime import datetime
import win32api
import config
import requests
import os

def initialize_api():
    """Initialize API base URL and headers."""
    base_url = "http://na.myconnectwise.net/v4_6_release/apis/3.0"
    client_id = config.client_id
    login_company = config.login_company
    public_key = config.public_key
    private_key = config.private_key
    auth = (f"{login_company}+{public_key}", private_key)
    headers = {
        "clientID": client_id,
        "Authorization": requests.auth._basic_auth_str(auth[0], auth[1])
    }
    return base_url, headers

def fetch_ticket_details(base_url, headers, ticket_id):
    """Fetch product details for a given Product ID."""
    if not ticket_id:
        return {}
    url = f"{base_url}/service/tickets/{ticket_id}?fields=id,company/name"
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        try:
            projectUrl = f"{base_url}/project/tickets/{ticket_id}?fields=id,company/name"
            response = requests.get(projectUrl, headers=headers)
            response.raise_for_status()
            return response.json()
        except requests.RequestException as projecte:
            return {"error": f"Service and Project fetch failed: {str(e)} | {str(projecte)}"}

def process_ticket(ticket_input, device_type, base_dir):
    """Process ticket and generate/print label."""
    ticket_number_input = ticket_input.replace("L", "").strip()
    base_url, headers = initialize_api()
    ticket_details = fetch_ticket_details(base_url, headers, ticket_number_input)
    
    if "error" in ticket_details:
        return False, f"API Error: {ticket_details['error']}"
    
    ticket_number = ticket_details.get("id")
    if not ticket_number:
        return False, "Invalid ticket number"
    
    date = datetime.now().strftime("%m/%d/%y")
    client_name = ticket_details.get("company", {}).get("name", "Unknown")
    
    if len(client_name) > 30:
        split_client_name = client_name[:30].split(" ")
        client_name = " ".join(split_client_name[:-1])
    
    output_path = os.path.join(base_dir, "ETS EQUIPMENT SHEET_Filled.docx")
    template_path = os.path.join(
        base_dir,
        f"ETS EQUIPMENT SHEET TEMPLATE - {device_type}.docx"
    )
    
    try:
        doc = Document(template_path)
    except FileNotFoundError:
        return False, f"Template not found: {template_path}"
    
    for paragraph in doc.paragraphs:
        if device_type == "Laptop":
            if "Ticket#" in paragraph.text:
                run = paragraph.add_run(f"{ticket_number}")
                run.bold = True
                run.font.size = Pt(11)
            elif "Date:" in paragraph.text:
                run = paragraph.add_run(f"{date}")
                run.font.size = Pt(11)
            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {client_name}")
                run.bold = True
                run.font.size = Pt(11)
        else:  # Normal
            if "Ticket #:" in paragraph.text:
                run = paragraph.add_run(f" {ticket_number}")
                run.bold = True
                run.font.size = Pt(11)
                run = paragraph.add_run(f" \tDate: {date}         Tech:______")
                run.font.size = Pt(11)
            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {client_name}")
                run.bold = True
                run.font.size = Pt(11)
    
    try:
        doc.save(output_path)
    except PermissionError:
        return False, "Please close the document and try again"
    
    try:
        win32api.ShellExecute(0, "printto", output_path, f'"{config.printer_name}"', ".", 0)
        return True, "Ticket processed and printed successfully"
    except Exception as e:
        return False, f"Printing failed: {str(e)}"