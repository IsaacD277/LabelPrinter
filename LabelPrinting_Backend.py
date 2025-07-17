from docx import Document
from docx.shared import Pt
from datetime import datetime
import win32api
import requests
import os

def initialize(config):
    # Initialize API base URL and headers
    baseUrl = config['CONNECTWISE_BASE_URL']
    clientId = config['CLIENT_ID']
    loginCompany = config['LOGIN_COMPANY']
    publicKey = config['PUBLIC_KEY']
    privateKey = config['PRIVATE_KEY']

    # Build authentication
    auth = (f"{loginCompany}+{publicKey}", privateKey)
    headers = {
        "clientID": clientId,
        "Authorization": requests.auth._basic_auth_str(auth[0], auth[1])
    }

    # Initialize base directory for templates
    baseDir = config['BASE_DIR']

    return baseUrl, headers, baseDir

def fetch_client(baseUrl, headers, ticketId):
    if not ticketId:
        return {}

    # Make API call to fetch service ticket details
    url = f"{baseUrl}/service/tickets/{ticketId}?fields=id,company/name"
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        details = response.json()
    except requests.RequestException as e:
        # If the ticket is not found, try fetching from project tickets
        try:
            projectUrl = f"{baseUrl}/project/tickets/{ticketId}?fields=id,company/name"
            response = requests.get(projectUrl, headers=headers)
            response.raise_for_status()
            details = response.json()
        except requests.RequestException as projecte:
            return {"error": f"Service and Project fetch failed: {str(e)} | {str(projecte)}"}

    # Extract client name from the response
    clientName = details.get("company").get("name")

    # If client name is longer than 30 characters, remove the end of the name on a word break
    if len(clientName) > 30:
        splitClientName = clientName[:30].split(" ")
        clientName = splitClientName[0]
        for word in splitClientName[1:-1]:
            clientName += " " + word

    return clientName

def update_Word_document(baseDir, ticketNumber, clientName, laptop = False):
    outputPath = os.path.join(baseDir, os.environ.get('TEMPLATE_FILLED'))
    date = datetime.now().strftime("%m/%d/%y")

    # Update the Laptop Word document with details for laptops
    if laptop:
        templatePath = os.path.join(baseDir, os.environ.get("TEMPLATE_LAPTOP"))
        doc = Document(templatePath)

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if "Ticket#" in paragraph.text:
                run = paragraph.add_run(f"{ticketNumber}")
                run.bold = True
                run.font.size = Pt(11)

            elif "Date:" in paragraph.text:
                run = paragraph.add_run(f"{date}")
                run.font.size = Pt(11)

            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {clientName}")
                run.bold = True
                run.font.size = Pt(11)

        try:
            doc.save(outputPath)
        except PermissionError:
            input(f"Please close the document and try again.")

    # Print the Word document with details for "Normal" boxes
    else:
        templatePath = os.path.join(baseDir, os.environ.get('TEMPLATE_NORMAL'))
        doc = Document(templatePath)

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if "Ticket #:" in paragraph.text:
                run = paragraph.add_run(f" {ticketNumber}")
                run.bold = True
                run.font.size = Pt(11)

                run = paragraph.add_run(f" \tDate: {date}         Tech:______")
                run.font.size = Pt(11)

            elif "Client Name:" in paragraph.text:
                run = paragraph.add_run(f" {clientName}")
                run.bold = True
                run.font.size = Pt(11)

        # Save the modified document
        try:
            doc.save(outputPath)
        except PermissionError:
            input(f"Please close the document and try again.")

    return outputPath

def print_label(filePath):
    try:
        win32api.ShellExecute(0, "printto", filePath, f'"{os.environ.get('PRINTER_NAME')}"', ".", 0)
        return True, "Ticket processed and printed successfully"
    except Exception as e:
        return False, f"Printing failed: {str(e)}"
    
def process_ticket(baseUrl, headers, baseDir, ticketId, laptop=False):
    # Fetch client name from the service ticket
    clientName = fetch_client(baseUrl, headers, ticketId)

    # Update Word document with ticket details
    outputPath = update_Word_document(baseDir, ticketId, clientName, laptop)

    # Print to label printer
    success, message = print_label(outputPath)

    if success:
        return True, f"Ticket #{ticketId} printed successfully"
    else:
        return False, f"Attempted to print ticket #{ticketId}. {message}"